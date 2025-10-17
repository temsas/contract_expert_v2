import PyPDF2
from docx import Document
import os
import logging

logger = logging.getLogger(__name__)


class FileProcessor:
    @staticmethod
    def extract_text_from_file(file_path, filename):
        """Извлекает текст из файла в зависимости от формата"""
        file_ext = os.path.splitext(filename)[1].lower()

        if file_ext == '.pdf':
            return FileProcessor._extract_from_pdf(file_path)
        elif file_ext in ['.doc', '.docx']:
            return FileProcessor._extract_from_docx(file_path)
        else:
            raise ValueError(f"Неподдерживаемый формат файла: {file_ext}")

    @staticmethod
    def _extract_from_pdf(file_path):
        """Извлекает текст из PDF"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""

                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += f"--- Страница {page_num + 1} ---\n{page_text}\n\n"

                if not text.strip():
                    raise Exception("PDF файл не содержит извлекаемого текста")

                logger.info(f"Извлечено {len(text)} символов из PDF")
                return text

        except Exception as e:
            logger.error(f"Ошибка при чтении PDF: {str(e)}")
            raise Exception(f"Ошибка при чтении PDF: {str(e)}")

    @staticmethod
    def _extract_from_docx(file_path):
        """Извлекает текст из DOC/DOCX"""
        try:
            doc = Document(file_path)
            text = ""

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " | "
                    text += "\n"
                text += "\n"

            if not text.strip():
                raise Exception("DOCX файл не содержит текста")

            logger.info(f"Извлечено {len(text)} символов из DOCX")
            return text

        except Exception as e:
            logger.error(f"Ошибка при чтении DOCX: {str(e)}")
            raise Exception(f"Ошибка при чтении DOCX: {str(e)}")